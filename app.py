"""
Read Smart AI - Persistent Chroma version (cleaned)
- Persistent vectorstore saved to ./chroma_db/
- Gemini embeddings via google.generativeai
- Uses LangChain RecursiveCharacterTextSplitter
- Streaming-safe, batched embeddings, better error handling
"""

import os
import io
from datetime import datetime
from textwrap import dedent
from typing import List, Dict, Tuple

import streamlit as st
from PyPDF2 import PdfReader

# Google Generative AI
import google.generativeai as genai

# LangChain text splitter (correct import path)
from langchain.text_splitter import RecursiveCharacterTextSplitter

# LangChain community Chroma wrapper (persistent)
from langchain_community.vectorstores import Chroma

# Optional export deps
try:
    from docx import Document as DocxDocument  # python-docx
    HAS_DOCX = True
except Exception:
    HAS_DOCX = False

try:
    from fpdf import FPDF
    HAS_FPDF = True
except Exception:
    HAS_FPDF = False


# -----------------------
# Config / Constants
# -----------------------
PERSIST_DIR = os.getenv("CHROMA_PERSIST_DIR", "./chroma_db")
CHROMA_COLLECTION_NAME = os.getenv("CHROMA_COLLECTION_NAME", "readsmart")
EMBED_MODEL = "models/text-embedding-004"
GEMINI_MODEL = "gemini-1.5-flash"
EMBED_BATCH_SIZE = 20


# -----------------------
# Utility: Init GenAI
# -----------------------
def init_genai() -> bool:
    """
    Configure the google.generativeai client from Streamlit secrets (or environment).
    Returns True if configured, False otherwise.
    """
    # Try Streamlit secrets first
    key = None
    try:
        if st.secrets:
            key = st.secrets.get("GOOGLE_API_KEY", None)
    except Exception:
        key = None

    # fall back to environment variable
    if not key:
        key = os.environ.get("GOOGLE_API_KEY", None)

    if not key:
        st.error("Missing Google API key. Set GOOGLE_API_KEY in Streamlit Secrets or environment.")
        return False

    genai.configure(api_key=key)
    return True


# -----------------------
# PDF -> text
# -----------------------
def pdf_to_text(file) -> List[Dict]:
    """
    Extracts text from PDF pages. Returns list of dicts: {"page": int, "text": str}
    """
    reader = PdfReader(file)
    pages = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            pages.append({"page": i + 1, "text": text})
    return pages


# -----------------------
# Text splitting / chunking
# -----------------------
def build_chunks_from_docs(docs_texts: List[Dict], chunk_size: int = 1000, chunk_overlap: int = 200) -> Tuple[List[str], List[Dict]]:
    """
    Accepts a list of documents in the form {"name": ..., "page": ..., "text": ...}
    Returns (chunks, metadata)
    """
    splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    chunks: List[str] = []
    metadata: List[Dict] = []

    for d in docs_texts:
        content = d.get("text", "")
        if not content.strip():
            continue
        splits = splitter.split_text(content)
        for idx, s in enumerate(splits):
            chunks.append(s)
            metadata.append({
                "source_name": d.get("name", "uploaded_pdf"),
                "page": d.get("page"),
                "chunk_index": idx
            })

    return chunks, metadata


# -----------------------
# Embedding helper (wraps genai.embed_content)
# -----------------------
def embed_texts_genai(texts: List[str]) -> List[List[float]]:
    """
    Given a list of texts, returns a list of embedding vectors using Gemini embedding model.
    Batched to EMBED_BATCH_SIZE within this function is not required because this function will
    be called by Chroma in batches — but we still handle a safe internal batch in case.
    """
    if not init_genai():
        raise RuntimeError("Gemini not configured")

    vectors: List[List[float]] = []
    for i in range(0, len(texts), EMBED_BATCH_SIZE):
        batch = texts[i:i + EMBED_BATCH_SIZE]
        try:
            resp = genai.embed_content(model=EMBED_MODEL, content=batch)
            # Some API variants return resp["embedding"] as a list of vectors
            vecs = resp.get("embedding", None) if isinstance(resp, dict) else None
            if vecs is None:
                # Attempt fallback attributes (some SDKs return different shapes)
                vecs = resp
            vectors.extend(vecs)
        except Exception as e:
            st.error(f"Embedding API error: {e}")
            raise e
    return vectors


# -----------------------
# Vectorstore Build (persistent)
# -----------------------
def build_vectorstore(chunks: List[str], metadata: List[Dict], persist_directory: str = PERSIST_DIR) -> Chroma:
    """
    Builds or updates a persistent Chroma vectorstore using an embedding function wrapper.
    This uses Chroma.from_texts with an embedding function that calls Gemini.
    """
    if not chunks:
        raise ValueError("No chunks to index")

    if not init_genai():
        st.stop()

    # embedding function wrapper expected by LangChain/Chroma
    def embedding_function(texts: List[str]) -> List[List[float]]:
        return embed_texts_genai(texts)

    # Ensure persist dir exists
    os.makedirs(persist_directory, exist_ok=True)

    # Create or load the chroma collection using from_texts.
    # NOTE: from_texts will call our embedding_function to generate vectors and persist them.
    try:
        chroma_db = Chroma.from_texts(
            texts=chunks,
            embedding=embedding_function,           # some LangChain wrappers use `embedding` param
            metadatas=metadata,
            collection_name=CHROMA_COLLECTION_NAME,
            persist_directory=persist_directory
        )
    except TypeError:
        # Some versions expect `embedding_function` as the kwarg name
        chroma_db = Chroma.from_texts(
            texts=chunks,
            embedding_function=embedding_function,
            metadatas=metadata,
            collection_name=CHROMA_COLLECTION_NAME,
            persist_directory=persist_directory
        )

    # Persist to disk (explicit if the wrapper supports persist())
    try:
        chroma_db.persist()
    except Exception:
        # not all wrappers expose persist(); ignore safely
        pass

    return chroma_db


# -----------------------
# Retrieve docs (semantic search)
# -----------------------
def retrieve_top_docs(vectorstore: Chroma, query: str, k: int = 4) -> List:
    """
    Retrieves top-k similar docs using an embedding for the query.
    """
    if not init_genai():
        st.stop()

    # get query vector
    resp = genai.embed_content(model=EMBED_MODEL, content=[query])
    q_vecs = resp.get("embedding") if isinstance(resp, dict) else resp
    if not q_vecs:
        raise RuntimeError("Unable to compute query embedding")
    query_vec = q_vecs[0]

    # Many Chroma wrappers expect list vectors; some have `similarity_search_by_vector`
    try:
        docs = vectorstore.similarity_search_by_vector(query_vec, k=k)
    except AttributeError:
        # try alternate name
        docs = vectorstore.similarity_search_with_score(query_vec, k=k)

    return docs


# -----------------------
# Gemini LLM call (simple)
# -----------------------
def call_gemini_with_context(prompt: str) -> str:
    """
    Calls the Gemini LLM with the provided prompt. Returns the assistant text.
    """
    if not init_genai():
        st.stop()

    try:
        resp = genai.generate(model=GEMINI_MODEL, prompt=prompt)
        # most SDKs return a dict-like with 'candidates' -> list -> content -> text
        if hasattr(resp, "candidates") or (isinstance(resp, dict) and "candidates" in resp):
            # support both attribute and dict access
            candidates = resp.candidates if hasattr(resp, "candidates") else resp["candidates"]
            if len(candidates) > 0:
                # some SDK content structure
                c = candidates[0]
                # try to pull content
                content = getattr(c, "content", None) or c.get("content", None) if isinstance(c, dict) else None
                if isinstance(content, list) and content and isinstance(content[0], dict):
                    # nested shape: content[0]["text"]
                    return content[0].get("text", str(resp))
                # try direct text attr
                text = getattr(c, "text", None) or c.get("text", None) if isinstance(c, dict) else None
                if text:
                    return text
        # Fallback: string-cast
        return str(resp)
    except Exception as e:
        st.error(f"LLM error: {e}")
        # attempt a simple fallback generation (best-effort)
        try:
            resp = genai.generate(model=GEMINI_MODEL, prompt=prompt)
            return str(resp)
        except Exception as e2:
            st.error(f"LLM fallback failed: {e2}")
            return "Error: failed to get response from Gemini."


# -----------------------
# Export helpers
# -----------------------
def export_chat_to_markdown(chat_history: List[Dict]) -> bytes:
    md_lines = []
    for msg in chat_history:
        role = msg.get("role", "assistant")
        text = msg.get("text", "")
        prefix = "**User:**" if role == "user" else "**Assistant:**"
        md_lines.append(f"{prefix}\n\n{text}\n\n---\n")
    md_bytes = "\n".join(md_lines).encode("utf-8")
    return md_bytes


def export_chat_to_docx(chat_history: List[Dict]) -> bytes:
    if not HAS_DOCX:
        raise RuntimeError("python-docx not installed")
    doc = DocxDocument()
    for msg in chat_history:
        role = msg.get("role", "assistant")
        text = msg.get("text", "")
        doc.add_heading("User" if role == "user" else "Assistant", level=4)
        doc.add_paragraph(text)
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()


def export_chat_to_pdf(chat_history: List[Dict]) -> bytes:
    if not HAS_FPDF:
        raise RuntimeError("fpdf not installed")
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=11)
    for msg in chat_history:
        role = msg.get("role", "assistant")
        header = "User:" if role == "user" else "Assistant:"
        text = msg.get("text", "")
        pdf.multi_cell(0, 6, f"{header}\n{text}\n")
        pdf.ln(1)
    bio = io.BytesIO()
    pdf.output(bio)
    bio.seek(0)
    return bio.read()


# -----------------------
# Streamlit app state init
# -----------------------
if "vector_store" not in st.session_state:
    st.session_state.vector_store = None

if "chunks" not in st.session_state:
    st.session_state.chunks = []

if "metadatas" not in st.session_state:
    st.session_state.metadatas = []

if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

if "uploaded_docs" not in st.session_state:
    st.session_state.uploaded_docs = []


# -----------------------
# UI Setup
# -----------------------
st.set_page_config(layout="wide", page_title="Read Smart AI — Persistent")
st.title("📘 Read Smart AI — Persistent (Chroma)")

# Sidebar controls
with st.sidebar:
    st.header("Controls")

    uploaded_files = st.file_uploader("Upload PDFs", type="pdf", accept_multiple_files=True)

    st.subheader("Embedding / Chunking")
    chunk_size = st.number_input("Chunk size", min_value=256, max_value=4000, value=1000, step=128)
    chunk_overlap = st.number_input("Chunk overlap", min_value=0, max_value=1000, value=200, step=50)
    top_k = st.slider("Top K (retrieval)", min_value=1, max_value=8, value=4)

    st.subheader("Actions")
    build_index_btn = st.button("Build / Rebuild Index")
    summarize_btn = st.button("Summarize All Docs")
    clear_btn = st.button("Clear All")

    st.markdown("---")
    st.subheader("Export")
    export_md_btn = st.button("Export Chat (Markdown)")
    export_docx_btn = st.button("Export Chat (Docx)")
    export_pdf_btn = st.button("Export Chat (PDF)")

    st.markdown("---")
    st.info(f"Persistence: {PERSIST_DIR}")


# -----------------------
# Document loading
# -----------------------
if uploaded_files:
    pages = []
    for f in uploaded_files:
        try:
            extracted = pdf_to_text(f)
            for p in extracted:
                pages.append({"name": f.name, "page": p["page"], "text": p["text"]})
        except Exception as e:
            st.error(f"Error reading PDF '{getattr(f, 'name', 'uploaded')}': {e}")

    if pages:
        st.session_state.uploaded_docs = pages
        st.success(f"Loaded {len(pages)} pages from {len(uploaded_files)} file(s).")


# -----------------------
# Clear all
# -----------------------
if clear_btn:
    # Clear only session-state (do not delete persisted chroma folder automatically)
    st.session_state.vector_store = None
    st.session_state.chunks = []
    st.session_state.metadatas = []
    st.session_state.chat_history = []
    st.session_state.uploaded_docs = []
    st.success("Cleared in-memory session state. Persistent DB (./chroma_db) was not deleted.")


# -----------------------
# Build / Rebuild index
# -----------------------
if build_index_btn:
    if not st.session_state.uploaded_docs:
        st.warning("Upload PDFs first before building the index.")
    else:
        with st.spinner("Building index… This may take a minute for many pages."):
            try:
                chunks, metas = build_chunks_from_docs(
                    st.session_state.uploaded_docs,
                    chunk_size=chunk_size,
                    chunk_overlap=chunk_overlap
                )
                st.session_state.chunks = chunks
                st.session_state.metadatas = metas

                chroma_db = build_vectorstore(chunks, metas, persist_directory=PERSIST_DIR)
                st.session_state.vector_store = chroma_db

                st.success(f"Indexed {len(chunks)} chunks into persistent Chroma at {PERSIST_DIR}.")
            except Exception as e:
                st.error(f"Failed to build index: {e}")


# -----------------------
# Main layout (chat + docs)
# -----------------------
col1, col2 = st.columns([3, 1])

with col1:
    st.subheader("Conversation")

    # Display chat history
    for msg in st.session_state.chat_history:
        bg = "#eef3ff" if msg["role"] == "user" else "#f7f7f7"
        st.markdown(
            f"<div style='background:{bg}; padding:10px; border-radius:8px; margin:6px 0'>{msg['text']}</div>",
            unsafe_allow_html=True
        )

    query = st.text_input("Ask a question…")
    ask_btn = st.button("Ask")


with col2:
    st.subheader("Uploaded Documents")
    if st.session_state.uploaded_docs:
        # show unique file names with pages count
        files = {}
        for d in st.session_state.uploaded_docs:
            files.setdefault(d["name"], 0)
            files[d["name"]] += 1
        for name, count in files.items():
            st.write(f"- **{name}** — {count} pages")
    else:
        st.info("No documents loaded.")


# -----------------------
# Export handlers
# -----------------------
if export_md_btn:
    if not st.session_state.chat_history:
        st.warning("No chat to export.")
    else:
        md_bytes = export_chat_to_markdown(st.session_state.chat_history)
        now = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        st.download_button("Download Markdown", data=md_bytes, file_name=f"readsmart_chat_{now}.md", mime="text/markdown")

if export_docx_btn:
    if not HAS_DOCX:
        st.error("python-docx not installed. Install it to enable Docx export.")
    elif not st.session_state.chat_history:
        st.warning("No chat to export.")
    else:
        docx_bytes = export_chat_to_docx(st.session_state.chat_history)
        now = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        st.download_button("Download Docx", data=docx_bytes, file_name=f"readsmart_chat_{now}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

if export_pdf_btn:
    if not HAS_FPDF:
        st.error("fpdf not installed. Install it to enable PDF export.")
    elif not st.session_state.chat_history:
        st.warning("No chat to export.")
    else:
        pdf_bytes = export_chat_to_pdf(st.session_state.chat_history)
        now = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        st.download_button("Download PDF", data=pdf_bytes, file_name=f"readsmart_chat_{now}.pdf", mime="application/pdf")


# -----------------------
# Question handler
# -----------------------
if ask_btn:
    if not query.strip():
        st.warning("Enter a question.")
    elif not st.session_state.vector_store:
        st.warning("Build the index first (click 'Build / Rebuild Index' in the sidebar).")
    else:
        with st.spinner("Retrieving relevant documents…"):
            try:
                docs = retrieve_top_docs(st.session_state.vector_store, query, k=top_k)
            except Exception as e:
                st.error(f"Retrieval error: {e}")
                docs = []

        if not docs:
            st.warning("No relevant documents found.")
        else:
            # Extract short context list
            taken = []
            for d in docs:
                # expected shape: Document object w/ .page_content and .metadata
                page_content = getattr(d, "page_content", None) or d.get("page_content", "")
                meta = getattr(d, "metadata", None) or d.get("metadata", {})
                taken.append({
                    "source_name": meta.get("source_name", "uploaded_pdf"),
                    "page": meta.get("page"),
                    "chunk_index": meta.get("chunk_index"),
                    "text": page_content
                })

            context = "\n\n---\n\n".join([
                f"Source: {t['source_name']} (page {t['page']})\n\n{t['text']}"
                for t in taken
            ])

            prompt = dedent(f"""
            Answer the question using ONLY the context below. If the answer is not contained in the context, say "I don't know."

            CONTEXT:
            {context}

            QUESTION:
            {query}

            ANSWER:
            """)

            with st.spinner("Thinking with Gemini…"):
                answer = call_gemini_with_context(prompt)

            # Update chat history
            st.session_state.chat_history.append({"role": "user", "text": query})
            st.session_state.chat_history.append({"role": "assistant", "text": answer})

            # show the answer immediately
            st.markdown("**Answer:**")
            st.markdown(answer)

            # persist vectorstore if available
            try:
                if st.session_state.vector_store:
                    st.session_state.vector_store.persist()
            except Exception:
                pass

            # rerun to update UI with history
            st.experimental_rerun()
